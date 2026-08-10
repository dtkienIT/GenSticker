from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANIFEST_PATH = ROOT / "project-docs.json"
OUTPUT_DIR = ROOT / "originals"

PRIMARY = "5B3DF5"
PRIMARY_DARK = "4338CA"
PINK = "DB2777"
INK = "0F172A"
MUTED = "475569"
LIGHT = "EEF2FF"
LIGHT_ALT = "F8FAFC"
BORDER = "CBD5E1"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_header(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("TRANG ")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc: Document, item: dict, meta: dict) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color, before, after in (
        ("Title", 28, INK, 0, 12),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 17, INK, 18, 8),
        ("Heading 2", 12.5, PRIMARY_DARK, 12, 5),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(5)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.7))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    left.width = Inches(4.4)
    right.width = Inches(2.3)
    for cell in (left, right):
        set_cell_shading(cell, PRIMARY_DARK)
        set_cell_margins(cell, 55, 100, 55, 100)
    p = left.paragraphs[0]
    run = p.add_run("GENSTICKER  /  PROJECT DOCUMENTATION")
    run.font.name = "Aptos"
    run.font.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"GS-DOC-{item['number']}")
    run.font.name = "Aptos"
    run.font.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(WHITE)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.7))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    set_cell_margins(left, 40, 0, 0, 0)
    set_cell_margins(right, 40, 0, 0, 0)
    p = left.paragraphs[0]
    run = p.add_run(f"{meta['branch']} @ {meta['commit']}  •  {meta['verifiedAt']}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(right.paragraphs[0])

    doc.core_properties.title = f"{item['number']} – {item['title']}"
    doc.core_properties.subject = item["summary"]
    doc.core_properties.author = "GenSticker project team"
    doc.core_properties.last_modified_by = "GenSticker project team"
    doc.core_properties.keywords = "GenSticker, source documentation, kien_v5"
    doc.core_properties.category = "Project documentation"
    doc.core_properties.comments = (
        "Source-derived clean-room documentation. No sample-project, AI-generated, or user image content."
    )
    doc.core_properties.created = datetime(2026, 8, 9, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2026, 8, 9, tzinfo=timezone.utc)
    doc.core_properties.revision = 1


def add_cover(doc: Document, item: dict, meta: dict) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    badge = doc.add_table(rows=1, cols=1)
    badge.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = badge.cell(0, 0)
    cell.width = Inches(1.05)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, 75, 130, 75, 130)
    p = cell.paragraphs[0]
    run = p.add_run(f"GS-DOC-{item['number']}")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(PRIMARY_DARK)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(18)
    p.add_run(item["title"])
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(item["subtitle"])

    accent = doc.add_table(rows=1, cols=2)
    accent.alignment = WD_TABLE_ALIGNMENT.LEFT
    accent.autofit = False
    accent.columns[0].width = Inches(4.8)
    accent.columns[1].width = Inches(1.7)
    left, right = accent.rows[0].cells
    set_cell_shading(left, PRIMARY)
    set_cell_shading(right, PINK)
    set_cell_margins(left, 18, 0, 18, 0)
    set_cell_margins(right, 18, 0, 18, 0)
    doc.add_paragraph()

    control = doc.add_table(rows=5, cols=2)
    control.style = "Table Grid"
    control.alignment = WD_TABLE_ALIGNMENT.LEFT
    control_values = [
        ("Phiên bản", meta["version"]),
        ("Baseline", f"{meta['branch']} @ {meta['commit']}"),
        ("Trạng thái", item["status"]),
        ("Ngày kiểm tra", meta["verifiedAt"]),
        ("Nguồn nội dung", "Source code hiện tại; mẫu chỉ dùng cấu trúc"),
    ]
    for index, (label, value) in enumerate(control_values):
        first, second = control.rows[index].cells
        set_cell_shading(first, LIGHT if index % 2 == 0 else LIGHT_ALT)
        set_cell_shading(second, WHITE if index % 2 == 0 else LIGHT_ALT)
        for cell in (first, second):
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        first.paragraphs[0].add_run(label).bold = True
        second.paragraphs[0].add_run(value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run(item["summary"])
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    toc = doc.add_paragraph()
    toc.paragraph_format.space_before = Pt(10)
    run = toc.add_run("NỘI DUNG")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(PRIMARY_DARK)
    for section in item.get("sections", []):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(section["title"])

    doc.add_page_break()


def add_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    set_repeat_header(p)
    run = p.add_run(title)
    run.font.color.rgb = RGBColor.from_string(INK)


def style_text_run(run, *, bold: bool = False, color: str = INK, size: float = 9.2) -> None:
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, PRIMARY_DARK)
        set_cell_border(cell, PRIMARY_DARK)
        set_cell_margins(cell, 80, 85, 80, 85)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_text_run(p.add_run(str(header)), bold=True, color=WHITE, size=8.6)
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row_values):
            cell = cells[col_index]
            set_cell_shading(cell, LIGHT_ALT if row_index % 2 else WHITE)
            set_cell_border(cell)
            set_cell_margins(cell, 70, 80, 70, 80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            style_text_run(p.add_run(str(value)), size=8.4 if len(headers) >= 4 else 9.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_section_content(doc: Document, section: dict) -> None:
    add_heading(doc, section["title"])
    section_type = section["type"]
    if section_type == "paragraphs":
        for text in section.get("paragraphs", []):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(7)
    elif section_type == "bullets":
        for text in section.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(text)
    elif section_type == "table":
        add_table(doc, section["headers"], section["rows"])
    elif section_type == "sources":
        p = doc.add_paragraph()
        run = p.add_run("Các đường dẫn dưới đây là nguồn kiểm chứng; số dòng có thể thay đổi khi code tiếp tục phát triển.")
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
        for source in section.get("sources", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.22)
            run = p.add_run(source)
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(PRIMARY_DARK)
    else:
        raise ValueError(f"Unsupported section type: {section_type}")


def add_figure(doc: Document, figure: dict, item: dict, figure_number: int) -> None:
    image_path = ROOT / figure["png"]
    if not image_path.exists():
        raise FileNotFoundError(f"Missing figure asset: {image_path}")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(2)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.65))
    shape._inline.docPr.set("title", figure["title"])
    shape._inline.docPr.set("descr", figure["alt"])

    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.add_run(
        f"Hình GS-DOC-{item['number']}.{figure_number} — {figure['title']}. "
        f"{figure['caption']}"
    )

    description = doc.add_table(rows=1, cols=1)
    description.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_properties = description.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    row_properties.append(cant_split)
    cell = description.cell(0, 0)
    set_cell_shading(cell, LIGHT_ALT)
    set_cell_border(cell, BORDER, "4")
    set_cell_margins(cell, 80, 100, 80, 100)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run("Mô tả sơ đồ: ")
    style_text_run(lead, bold=True, color=PRIMARY_DARK, size=8.3)
    body = p.add_run(figure["alt"])
    style_text_run(body, color=MUTED, size=8.3)

    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.space_after = Pt(8)
    run = source.add_run(
        "Nguồn: " + "; ".join(figure.get("sourceRefs", []))
        + f" · baseline {item.get('baselineLabel', 'kien_v5')}"
    )
    run.font.name = "Aptos Mono"
    run.font.size = Pt(7.4)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build_document(item: dict, meta: dict, figures: dict) -> Path:
    doc = Document()
    configure_document(doc, item, meta)
    add_cover(doc, item, meta)
    visuals_by_section: dict[str, list[dict]] = {}
    for visual in item.get("visuals", []):
        section = visual.get("afterSection")
        if section:
            visuals_by_section.setdefault(section, []).append(visual)
    figure_number = 1
    for section in item.get("sections", []):
        add_section_content(doc, section)
        for visual in visuals_by_section.get(section["title"], []):
            add_figure(doc, figures[visual["figureId"]], item, figure_number)
            figure_number += 1
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / item["filename"]
    doc.save(output_path)
    return output_path


def main() -> int:
    payload = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    meta = payload["meta"]
    figures = payload.get("figures", {})
    built = []
    for item in payload["documents"]:
        if item["kind"] == "DOCX":
            built.append(build_document(item, meta, figures))
    for path in built:
        print(path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
